import json
import os
from docx import Document
from openpyxl import Workbook

async def generate_word_document(title: str, content: str, filepath: str) -> str:
    """
    Generates a Word document with the given title and markdown-like content.
    For simplicity, we handle basic text paragraphs.
    """
    try:
        doc = Document()
        doc.add_heading(title, level=1)
        
        # Split content by newlines and add as paragraphs
        # A full implementation would parse markdown into bold/italic/headings
        for line in content.split('\n'):
            if line.strip():
                if line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                else:
                    doc.add_paragraph(line)
                    
        doc.save(filepath)
        return f"Successfully generated Word document at {os.path.abspath(filepath)}"
    except Exception as e:
        return f"Failed to generate Word document: {str(e)}"

import csv
import io

async def generate_excel_spreadsheet(csv_data: str, filepath: str) -> str:
    """
    Generates an Excel spreadsheet from a CSV string.
    """
    try:
        if not csv_data or not csv_data.strip():
            return "Error: csv_data cannot be empty."
            
        wb = Workbook()
        ws = wb.active
        
        # Read the CSV data
        reader = csv.reader(io.StringIO(csv_data.strip()))
        for row in reader:
            ws.append(row)
            
        wb.save(filepath)
        return f"Successfully generated Excel spreadsheet at {os.path.abspath(filepath)}"
    except Exception as e:
        return f"Failed to generate Excel spreadsheet: {str(e)}"
